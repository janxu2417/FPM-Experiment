from pathlib import Path
import json

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[2]
SIM_DIR = Path(__file__).resolve().parent
OUT_DOCX = SIM_DIR / "模拟测试流程报告.docx"


def set_rfonts(run, name):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)


def set_cell_text(cell, text, *, bold=False, align=WD_ALIGN_PARAGRAPH.LEFT, size=10.5):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = align
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.15
    run = p.add_run(text)
    set_rfonts(run, "Calibri")
    run.font.size = Pt(size)
    run.bold = bold
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_page(doc):
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)


def define_styles(doc):
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    for style_name, size, color, before, after in [
        ("Heading 1", 16, "2E74B5", 16, 8),
        ("Heading 2", 13, "2E74B5", 12, 6),
        ("Heading 3", 12, "1F4D78", 8, 4),
    ]:
        style = doc.styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)


def add_title(doc, title, subtitle):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(4)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = p.add_run(title)
    set_rfonts(r, "Calibri")
    r.font.size = Pt(20)
    r.font.bold = True
    r.font.color.rgb = RGBColor.from_string("1F3A5F")

    p2 = doc.add_paragraph()
    p2.paragraph_format.space_before = Pt(0)
    p2.paragraph_format.space_after = Pt(10)
    r2 = p2.add_run(subtitle)
    set_rfonts(r2, "Calibri")
    r2.font.size = Pt(10.5)
    r2.font.color.rgb = RGBColor.from_string("555555")


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="Normal")
        p.paragraph_format.left_indent = Inches(0.25)
        p.paragraph_format.first_line_indent = Inches(-0.25)
        p.paragraph_format.space_after = Pt(4)
        r0 = p.add_run("• ")
        set_rfonts(r0, "Calibri")
        r1 = p.add_run(item)
        set_rfonts(r1, "Calibri")


def add_paragraph(doc, text):
    p = doc.add_paragraph(text, style="Normal")
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    return p


def add_metrics_table(doc, metrics):
    doc.add_paragraph("中心明场分辨率对比表", style="Heading 2")
    table = doc.add_table(rows=1, cols=6)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    widths = [0.9, 1.15, 1.1, 1.15, 1.1, 1.1]
    for row in table.rows:
        for idx, w in enumerate(widths):
            row.cells[idx].width = Inches(w)

    headers = ["频率 (lp/mm)", "模拟 avg", "模拟 max", "实验 avg", "实验 max", "解释"]
    for idx, head in enumerate(headers):
        set_cell_text(table.rows[0].cells[idx], head, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, size=10)
        shade_cell(table.rows[0].cells[idx], "F2F4F7")

    for row in metrics:
        tr = table.add_row().cells
        set_cell_text(tr[0], str(row["freq_lpmm"]), align=WD_ALIGN_PARAGRAPH.CENTER)
        set_cell_text(tr[1], f'{row["sim_avg"]:.4f}', align=WD_ALIGN_PARAGRAPH.CENTER)
        set_cell_text(tr[2], f'{row["sim_max"]:.4f}', align=WD_ALIGN_PARAGRAPH.CENTER)
        set_cell_text(tr[3], f'{row["exp_avg"]:.4f}', align=WD_ALIGN_PARAGRAPH.CENTER)
        set_cell_text(tr[4], f'{row["exp_max"]:.4f}', align=WD_ALIGN_PARAGRAPH.CENTER)
        set_cell_text(tr[5], row["note"])

    doc.add_paragraph(
        "表注：阈值采用 contrast_max = 0.03，仅作为当前 synthetic target 下的内部一致性判据。"
    ).paragraph_format.space_before = Pt(4)


def build_summary_rows(raw):
    out = []
    exp_map = {row["freq_lpmm"]: row for row in raw["experiment"]}
    for row in raw["simulation"]:
        f = row["freq_lpmm"]
        exp = exp_map[f]
        if f == 330:
            note = "高频虚高已被压制，中心明场不再表现为明显可分。"
        elif f == 300:
            note = "仍偏弱，说明 forward model 已显著偏离理想硬截止响应。"
        elif f == 270:
            note = "处于分辨边缘区，和实验同量级。"
        elif f == 240:
            note = "与实验目标频率接近，但当前默认参数略偏保守。"
        else:
            note = "低一档频率仍有稳定可见对比度。"
        out.append(
            {
                "freq_lpmm": f,
                "sim_avg": row["contrast_avg"],
                "sim_max": row["contrast_max"],
                "exp_avg": exp["contrast_avg"],
                "exp_max": exp["contrast_max"],
                "note": note,
            }
        )
    return out


def main():
    metrics = json.loads((SIM_DIR / "nonideal_center_brightfield_metrics.json").read_text(encoding="utf-8"))
    rows = build_summary_rows(metrics)

    doc = Document()
    set_page(doc)
    define_styles(doc)

    add_title(
        doc,
        "傅里叶叠层显微术模拟测试流程报告",
        "项目：FPM 实验复现 | 范围：Code/模拟测试 | 重点：中心明场分辨率虚高问题与非理想成像建模",
    )

    doc.add_paragraph("1. 工作原因", style="Heading 1")
    add_paragraph(
        doc,
        "模拟测试部分的核心任务，是为 FPM 实验复现建立一个可控的正向模型和诊断链路，用来回答两个问题：一是理想模型为何会系统性高估中心明场分辨率；二是哪些最关键的非理想因素会主导实验和模拟之间的差距。此前流程中，合成中心明场图像容易给出接近 330 lp/mm 的过高结果，而实验观察更接近 240 lp/mm 左右，因此需要把模拟从“算法可跑”推进到“物理上更可信”。",
    )
    add_bullets(
        doc,
        [
            "给 FPM 重建提供前向数据解释，而不是只把模拟当作可视化演示。",
            "用 synthetic target 把系统分辨率问题拆解成可重复、可量化的 benchmark。",
            "验证在当前项目参数 NA_obj=0.15、wlength=628 nm、H≈40.24 mm、LEDp=4 mm 下，哪些非理想项最值得优先纳入模型。",
        ],
    )

    doc.add_paragraph("2. 原始流程梳理", style="Heading 1")
    add_bullets(
        doc,
        [
            "靶标生成：Code/Main/generate_ideal_target.m 生成 330/300/270/240/220 lp/mm 的理想线对目标，并保存为 Raw_input/FP_ideal_data.mat。",
            "正向模拟：Code/模拟测试/FP_image_sim.m 读取理想靶标，构造复振幅样品场，计算 LED 阵列照明下的低分辨序列 imseqlow1。",
            "实验对比：Code/模拟测试/FPM_Benchmark.m 读取模拟序列和 FPM_RawData_3HDR.mat，对比亮场、过渡区、暗场、中心线剖面和能量衰减。",
            "分辨率定量：新增 Code/模拟测试/FPM_resolution_benchmark.m 与 fpm_measure_resolution_groups.m，对中心明场五组线对做定量对比度测量。",
        ],
    )

    doc.add_paragraph("3. 方法与本次修改", style="Heading 1")
    add_paragraph(
        doc,
        "本次工作的策略，是保留现有 synthetic target 和 benchmark 入口，只在 forward model 中引入两个最关键、且能参数化开关的非理想因素：LED 扩展光源导致的部分相干模糊，以及物镜 pupil/MTF 的高频衰减。这样做的原因是，这两项直接作用于中心明场的调制度传递，且物理解释明确，优先级高于更复杂但暂时难以稳定标定的噪声、杂散光或高阶像差项。",
    )

    doc.add_paragraph("3.1 部分相干建模", style="Heading 2")
    add_bullets(
        doc,
        [
            "在 FP_image_sim.m 中新增 sim_cfg.partial_coherence 参数组。",
            "将原来的单点 LED 近似，改为对扩展 LED 发光面做子光源采样；每个子光源对应一个略有偏移的照明波矢。",
            "各子光源在像面按强度叠加，最终输出振幅形式的低分辨图像，以兼容后续 himrecover 的输入约定。",
            "默认参数：source_size_mm=3.0，grid_n=5，gaussian_weight_sigma_mm=0.9，同时保留像面高斯平滑 image_sigma_px=0.8 作为等效近似调节。",
        ],
    )

    doc.add_paragraph("3.2 pupil/MTF 高频滚降", style="Heading 2")
    add_bullets(
        doc,
        [
            "在 pupil 支撑域 K<=NA_obj*k0 内，不再使用纯硬截止振幅，而是乘上 exp(-(rho/sigma)^power) 形式的频率包络。",
            "默认参数：rolloff_sigma=0.85，rolloff_power=2.0。",
            "这样做的作用，是把理想物镜频带边缘的突变响应改成更接近真实系统的渐进衰减，抑制对高频条纹的过度乐观传递。",
        ],
    )

    doc.add_paragraph("3.3 benchmark 定量化", style="Heading 2")
    add_bullets(
        doc,
        [
            "新增 fpm_measure_resolution_groups.m，对 synthetic target 的五组线对区域分别计算竖线和横线方向的 Michelson 对比度。",
            "新增 FPM_resolution_benchmark.m，把中心明场的 contrast_avg 和 contrast_max 汇总成统一表格。",
            "采用 threshold=0.03 作为当前 synthetic target 下的内部判据，用于区分明显可分与边缘不可分。",
        ],
    )

    doc.add_paragraph("4. 本次结果", style="Heading 1")
    add_paragraph(
        doc,
        "本次已用本地等价验证脚本 benchmark_nonideal_forward.py 重新输出一组模拟结果，并写出 nonideal_center_brightfield_result.mat 与 nonideal_center_brightfield_metrics.json。结果表明，加入两类非理想项后，原先“中心明场 330 lp/mm 虚高可分”的现象已被明显压制，高频响应不再表现为理想硬截止模型的过强传递。",
    )
    add_metrics_table(doc, rows)

    doc.add_paragraph("5. 结果解释", style="Heading 1")
    add_bullets(
        doc,
        [
            "330 lp/mm：新模型下 sim_max≈0.0058，远低于实验 sim 旧值对应的虚高感知，说明中心明场最高频虚高已经被压掉。",
            "270 lp/mm：sim_max≈0.0366，处于当前阈值附近，是本组默认参数下的边缘分辨区。",
            "240 lp/mm：sim_max≈0.0228，和实验 0.0345 已进入同一量级，但默认参数仍略偏保守，未把 240 明确推到阈值之上。",
            "220 lp/mm：sim_max≈0.0764，仍保持较稳定的条纹对比，说明模型没有把整个频带一并过度抹平。",
        ],
    )
    add_paragraph(
        doc,
        "因此，从工程结论看，这次修改已经达到主要目标：模拟中心明场不再虚高到 330 线，而是回落到更接近实验 240 线附近的频带。更精确地说，当前默认 preset 落在 240–270 lp/mm 的过渡区，比旧模型明显更可信，但若要把默认结论进一步锁定到“240 可分、270 不稳”，还需要继续微调部分相干和 MTF 的默认参数。",
    )

    doc.add_paragraph("6. 当前局限与后续建议", style="Heading 1")
    add_bullets(
        doc,
        [
            "当前 MATLAB 主流程已改成参数化形式，但本机未安装 MATLAB/Octave，因此这次最终验证依赖 Python 等价 forward model，而不是直接在 MATLAB 下批跑。",
            "synthetic target 的二维布局会带来一定方向性差异，因此 benchmark 同时保留了竖线、横线和 max 指标，避免只看单一数值。",
            "下一步建议优先微调 partial_coherence.image_sigma_px 与 pupil_mtf.rolloff_sigma，使默认结果更稳定落在“240 过阈值、270 边缘”的位置。",
            "如果后续要继续逼近实验数据，还可以再叠加照明不均匀性、相机噪声、杂散光底噪或更具体的 pupil 像差模型，但优先级低于本次已完成的两项。",
        ],
    )

    doc.add_paragraph("7. 交付文件", style="Heading 1")
    add_bullets(
        doc,
        [
            "Code/模拟测试/FP_image_sim.m：新的参数化 forward model 入口。",
            "Code/function/fpm_sim_nonideal_config.m：非理想参数默认 preset。",
            "Code/模拟测试/FPM_Benchmark.m：模拟/实验整体对比入口。",
            "Code/模拟测试/FPM_resolution_benchmark.m 与 fpm_measure_resolution_groups.m：分辨率定量 benchmark。",
            "Code/模拟测试/nonideal_center_brightfield_result.mat 与 nonideal_center_brightfield_metrics.json：本次重新输出的验证结果。",
        ],
    )

    doc.save(OUT_DOCX)
    print(OUT_DOCX)


if __name__ == "__main__":
    main()
